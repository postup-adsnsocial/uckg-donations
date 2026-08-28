from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "dossier" / "Dossie_UCKG_Donations_Cliente.docx"
LOGO = ROOT / "apps" / "web" / "public" / "universal-logo.png"

NAVY = "003B66"
BLUE = "005A9C"
MID_BLUE = "1769AA"
LIGHT_BLUE = "EAF3F8"
PALE_BLUE = "F4F8FB"
INK = "172033"
MUTED = "5D6B82"
LINE = "D7E0E8"
WHITE = "FFFFFF"
GREEN = "1F6A4B"
AMBER = "8A6200"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, 8.5, MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16
    for name, size, color, before, after in [
        ("Heading 1", 17, NAVY, 18, 7),
        ("Heading 2", 13, BLUE, 13, 5),
        ("Heading 3", 11.5, NAVY, 9, 3),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("UCKG DONATIONS  |  Dossiê de Produto")
    set_run_font(run, 8.5, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_number(footer)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, 28, NAVY, bold=True)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(15)
        r = p.add_run(subtitle)
        set_run_font(r, 13, MUTED)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_run_font(r, 8.7, BLUE, bold=True)
    r.font.all_caps = True


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, 10.3, INK, bold=True)
    r = p.add_run(text)
    set_run_font(r, 10.3, INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.17)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, 10.2, INK)


def restart_list_numbering(doc, paragraphs):
    numbering = doc.part.numbering_part.element
    list_style = doc.styles["List Number"]._element
    style_num_id = list_style.pPr.numPr.numId.val
    base_num = next(
        num for num in numbering.num_lst if num.numId == style_num_id
    )
    abstract_num_id = base_num.abstractNumId.val
    next_num_id = max((num.numId for num in numbering.num_lst), default=0) + 1

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    number.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    number.append(override)
    numbering.append(number)

    for paragraph in paragraphs:
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = num_pr.find(qn("w:ilvl"))
        if ilvl is None:
            ilvl = OxmlElement("w:ilvl")
            num_pr.append(ilvl)
        ilvl.set(qn("w:val"), "0")
        num_id = num_pr.find(qn("w:numId"))
        if num_id is None:
            num_id = OxmlElement("w:numId")
            num_pr.append(num_id)
        num_id.set(qn("w:val"), str(next_num_id))


def add_numbered(doc, items):
    paragraphs = []
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, 10.2, INK)
        paragraphs.append(p)
    restart_list_numbering(doc, paragraphs)


def add_callout(doc, title, body, tone="blue"):
    color = LIGHT_BLUE if tone == "blue" else "F8F4E8"
    title_color = BLUE if tone == "blue" else AMBER
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, 10.5, title_color, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(body)
    set_run_font(r, 9.8, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, 9.2, WHITE, bold=True)
    for index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            if index % 2 == 1:
                set_cell_shading(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, 9.1, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metadata(doc, entries):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2500, 6860])
    for label, value in entries:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_BLUE)
        for cell in cells:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
        r = cells[0].paragraphs[0].add_run(label)
        set_run_font(r, 9.2, NAVY, bold=True)
        r = cells[1].paragraphs[0].add_run(value)
        set_run_font(r, 9.2, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    configure(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(0.68))
    add_kicker(doc, "Plataforma corporativa de gestão de contribuições")
    add_title(doc, "UCKG Donations", "Dossiê funcional e técnico para apresentação ao cliente")
    add_body(doc, "Uma plataforma web multi-igreja para registrar, proteger, organizar e acompanhar contribuições, membros e relatórios financeiros em uma operação centralizada.")
    add_callout(doc, "Proposta de valor", "A plataforma oferece visibilidade e padronização para a operação financeira de múltiplas igrejas, preservando o isolamento de dados de cada igreja e a rastreabilidade de cada lançamento.")
    add_metadata(doc, [
        ("Produto", "UCKG Donations - Gestão Financeira Multi-Igreja"),
        ("Público", "Administração central, administradores de igreja, operadores financeiros e auditores"),
        ("Mercado-alvo", "Rede de igrejas nos Estados Unidos, com capacidade planejada para até 150 unidades"),
        ("Idiomas", "Português do Brasil, inglês e espanhol"),
        ("Documento", "Versão 1.0 | 4 de agosto de 2026"),
        ("Classificação", "Material de apresentação comercial e especificação de escopo"),
    ])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENCIAL - Preparado para avaliação da solução")
    set_run_font(r, 8.4, MUTED, bold=True)

    page_break(doc)
    add_kicker(doc, "01 | Visão executiva")
    doc.add_heading("Resumo da solução", level=1)
    add_body(doc, "O UCKG Donations centraliza o registro e a consulta de contribuições recebidas pelas igrejas, com identificação opcional do membro, imagem privada do envelope, histórico financeiro e relatórios em PDF. A solução foi desenhada para uma administração em rede, onde cada unidade opera em seu próprio ambiente lógico, sem acesso aos dados das demais.")
    add_body(doc, "A experiência operacional é web, responsiva e localizada para três idiomas. O objetivo é substituir processos dispersos de planilhas, e-mails e arquivos locais por um fluxo único, verificável e acessível a usuários autorizados.")
    doc.add_heading("Resultados esperados", level=2)
    add_bullets(doc, [
        "Padronizar o registro de contribuições e a documentação de envelopes entre as igrejas.",
        "Reduzir retrabalho na consolidação de informações financeiras e na preparação de relatórios.",
        "Restringir o acesso a informações de membros e doações conforme a função de cada operador.",
        "Permitir que a administração central acompanhe a estrutura de igrejas sem expor dados entre unidades.",
        "Disponibilizar relatórios consistentes e arquivados, com download seguro quando necessário.",
    ])
    doc.add_heading("Princípios do produto", level=2)
    add_matrix(doc, ["Princípio", "Aplicação prática"], [
        ("Multi-igreja por padrão", "Todo dado operacional pertence a uma igreja específica; o contexto de trabalho é escolhido explicitamente."),
        ("Segurança por função", "Ações de leitura, lançamento e gestão são liberadas apenas a perfis autorizados."),
        ("Rastreabilidade", "Cada contribuição guarda data, valor, forma de recebimento, membro opcional, observação e operador responsável."),
        ("Experiência inclusiva", "A mesma navegação é oferecida em português, inglês e espanhol, em desktop e dispositivos móveis."),
    ], [2600, 6760])

    doc.add_heading("Fluxo operacional em alto nível", level=2)
    add_numbered(doc, [
        "O operador entra com credenciais administrativas protegidas.",
        "Seleciona a igreja na qual possui autorização para trabalhar.",
        "Consulta ou cadastra membros e registra contribuições identificadas ou anônimas.",
        "Anexa a imagem do envelope, quando aplicável, em armazenamento privado.",
        "Consulta o histórico por período ou membro e gera relatórios PDF para análise e arquivo.",
    ])

    page_break(doc)
    add_kicker(doc, "02 | Perfis, acesso e governança")
    doc.add_heading("Perfis de usuário", level=1)
    add_body(doc, "A plataforma separa a identidade administrativa do cadastro de membros. Um membro da igreja não recebe acesso ao sistema por estar cadastrado; somente usuários administrativos, com uma função atribuída, podem entrar.")
    add_matrix(doc, ["Perfil", "Responsabilidades principais", "Acesso"], [
        ("Administrador da plataforma", "Administração central da rede e gestão de igrejas ativas.", "Acesso explícito às igrejas ativas e às funções de administração central."),
        ("Administrador da igreja", "Gestão local de cadastros, contribuições e configurações permitidas.", "Leitura e escrita na igreja de sua responsabilidade."),
        ("Operador financeiro", "Lançamento e consulta de contribuições; consulta de membros conforme necessidade operacional.", "Operações financeiras na igreja autorizada."),
        ("Auditor", "Acompanhamento autorizado e consulta de informações administrativas permitidas.", "Acesso de leitura limitado, sem lançamento financeiro."),
    ], [1900, 4300, 3160])
    doc.add_heading("Regras de acesso", level=2)
    add_bullets(doc, [
        "A igreja ativa é informada em todas as operações de domínio; requisições sem contexto são negadas.",
        "O sistema valida a associação entre usuário e igreja antes de consultar ou alterar qualquer registro.",
        "Tentativas de acessar dados de outra igreja são bloqueadas sem revelar informações da unidade consultada.",
        "A administração central pode criar, editar e arquivar igrejas de forma controlada; o último cadastro ativo não pode ser arquivado por engano.",
        "Os registros financeiros são vinculados ao operador que os criou, preservando a origem operacional do lançamento.",
    ])
    add_callout(doc, "Governança recomendada", "A definição de quais pessoas serão administradores, operadores financeiros e auditores deve ser aprovada pela organização. A administração central deve revisar regularmente os acessos ativos, especialmente em caso de transferência de equipe ou desligamento.", tone="amber")

    page_break(doc)
    add_kicker(doc, "03 | Funcionalidades de cadastro e igrejas")
    doc.add_heading("Gestão de igrejas", level=1)
    add_body(doc, "A administração central mantém um cadastro de igrejas que serve de base para a segregação de dados e a operação da rede. Cada igreja possui um identificador técnico interno, nome, status, idioma, fuso horário e dados de contato e endereço quando configurados.")
    add_matrix(doc, ["Função", "Descrição", "Resultado"], [
        ("Cadastrar igreja", "Criação de uma nova unidade com nome informado pela administração.", "A plataforma gera um identificador técnico único e aplica configurações iniciais da operação nos EUA."),
        ("Editar igreja", "Atualização controlada do nome e dos dados administrativos permitidos.", "Mantém a identidade técnica da unidade e evita troca acidental de propriedade dos registros."),
        ("Arquivar igreja", "Desativação segura de uma unidade que não deve mais operar na plataforma.", "A igreja deixa de estar disponível para novas operações, preservando o histórico necessário."),
        ("Selecionar igreja ativa", "Escolha explícita da unidade em que o operador está trabalhando.", "Todas as consultas e lançamentos seguintes são filtrados por essa igreja."),
    ], [1800, 4300, 3260])
    doc.add_heading("Cadastro de membros", level=1)
    add_body(doc, "O módulo de membros organiza a base cadastral local de cada igreja. O cadastro suporta nome completo, e-mail, telefone, endereço nos Estados Unidos, observações e situação ativo/inativo.")
    add_bullets(doc, [
        "Listagem por igreja com busca por nome, e-mail ou telefone.",
        "Paginação e ordenação estável para manter a consulta previsível em bases maiores.",
        "Cadastro, visualização de detalhes, edição e desativação segura, sem exclusão física do histórico.",
        "Validações para e-mail normalizado, telefone em padrão internacional e CEP dos Estados Unidos.",
        "Proteção contra duplicidade de e-mail dentro da mesma igreja, sem comparar dados de outras unidades.",
        "Consulta do histórico de contribuições vinculado ao perfil do membro quando aplicável.",
    ])

    page_break(doc)
    add_kicker(doc, "04 | Funcionalidades financeiras")
    doc.add_heading("Registro de contribuições e envelopes", level=1)
    add_body(doc, "Cada contribuição é registrada como um lançamento financeiro positivo, associado à igreja ativa. O lançamento pode estar ligado a um membro ou ser registrado como anônimo, sem a necessidade de criar um cadastro artificial.")
    add_matrix(doc, ["Campo", "Especificação"], [
        ("Valor", "Armazenado em centavos de dólar para preservar precisão nos cálculos e relatórios."),
        ("Data de recebimento", "Data operacional informada no lançamento para permitir filtros e consolidações por período."),
        ("Membro", "Vínculo opcional; o sistema valida que o membro pertence à mesma igreja do lançamento."),
        ("Forma de recebimento", "Dinheiro, cartão ou cheque, com identificação no histórico e nos relatórios."),
        ("Observações", "Campo opcional para registrar contexto operacional do lançamento."),
        ("Operador", "Usuário administrativo responsável pela criação do registro."),
        ("Imagem do envelope", "Arquivo JPEG ou PNG de até 4 MB, ligado ao lançamento e armazenado de forma privada."),
    ], [2450, 6910])
    doc.add_heading("Histórico e consulta", level=2)
    add_bullets(doc, [
        "Lista cronológica de contribuições, com valor, data, membro ou anonimato, forma de recebimento e indicação de imagem anexada.",
        "Filtros por período e membro para localizar lançamentos específicos.",
        "Tela de detalhes com valor, igreja, data, operador, método de recebimento, observações e documento de envelope.",
        "Visualização segura da imagem do envelope somente para usuários com a permissão correspondente.",
        "Resumo de quantidade e total para os filtros utilizados na consulta.",
    ])
    add_callout(doc, "Tratamento do arquivo", "A imagem do envelope não fica pública na internet. Em produção, o arquivo é enviado a um bucket privado e é recuperado por um link temporário controlado ou por resposta autenticada da aplicação.")

    page_break(doc)
    add_kicker(doc, "05 | Painéis e relatórios")
    doc.add_heading("Painel operacional", level=1)
    add_body(doc, "O painel inicial mostra o contexto da igreja selecionada e direciona o operador aos módulos de membros, envelopes e relatórios. Os dados exibidos são carregados a partir dos registros reais da igreja ativa, sem misturar informações de outras igrejas.")
    doc.add_heading("Relatórios financeiros", level=1)
    add_body(doc, "O módulo de relatórios consolida as contribuições de um período definido e permite gerar um arquivo PDF privado. O operador pode trabalhar com períodos comuns, como mês atual, mês anterior, últimos 30 dias e ano atual, ou informar uma data inicial e final personalizada.")
    add_matrix(doc, ["Tipo de relatório", "Conteúdo"], [
        ("Detalhado", "Lista as contribuições do período, incluindo membro ou anonimato, data, método, valor, operador e observações. Pode incluir as imagens de envelope."),
        ("Totais por membro", "Agrupa quantidade e valor total por membro, preservando os filtros e a separação por igreja."),
        ("Totais por forma de recebimento", "Agrupa quantidade e valor total por dinheiro, cartão e cheque."),
        ("Arquivo histórico", "Os PDFs gerados ficam registrados por igreja para consulta e novo download pelos usuários autorizados."),
    ], [2600, 6760])
    doc.add_heading("Características do PDF", level=2)
    add_bullets(doc, [
        "Identificação da igreja, período consultado e total consolidado.",
        "Formatação preparada para leitura, impressão e compartilhamento interno.",
        "Inclusão opcional de imagens dos envelopes no relatório detalhado.",
        "Armazenamento privado e download autenticado, com possibilidade de URL temporária em ambiente gerenciado.",
    ])

    page_break(doc)
    add_kicker(doc, "06 | Segurança, privacidade e arquitetura")
    doc.add_heading("Proteção de acesso e dados", level=1)
    add_matrix(doc, ["Camada", "Especificação de proteção"], [
        ("Autenticação", "Senhas protegidas por hash com salt e sessões opacas. O token de sessão é armazenado somente em forma de hash no banco."),
        ("Sessão", "Cookie HTTP-only, SameSite=Strict e Secure em produção; validade operacional de 12 horas."),
        ("Autorização", "Permissões declaradas por rota e por perfil. A plataforma adota negação por padrão para rotas sem autorização explícita."),
        ("Isolamento multi-igreja", "Registros possuem vínculo com a igreja. Consultas, inserções e relacionamentos são verificados no contexto do tenant ativo, na API e no banco."),
        ("Banco de dados", "PostgreSQL com restrições, chaves estrangeiras, índices e políticas de segurança que reforçam a segregação de dados."),
        ("Arquivos privados", "Imagens de envelope e PDFs ficam fora das tabelas relacionais, em armazenamento privado, com metadados e vínculo à igreja."),
        ("Transporte", "Implantação recomendada com HTTPS, cabeçalhos de segurança, CORS controlado e limites de tamanho para requisições."),
    ], [2450, 6910])
    doc.add_heading("Como a plataforma reduz o risco de vazamento", level=2)
    add_numbered(doc, [
        "O navegador nunca recebe a credencial de serviço do armazenamento ou do banco. As operações sensíveis acontecem na API, em ambiente controlado.",
        "O usuário entra em uma sessão HTTP-only: scripts executados no navegador não conseguem ler o token de sessão diretamente.",
        "Cada operação exige uma igreja ativa autorizada. A API valida a associação do usuário com a igreja antes de executar a ação solicitada.",
        "O banco reforça o isolamento por Row-Level Security (RLS) e por contexto transacional. Mesmo uma consulta que não aplique o filtro correto não deve retornar dados de outra igreja para a credencial de execução da aplicação.",
        "A associação entre contribuição, membro, arquivo e igreja é validada por chaves estrangeiras e por verificações de mesma igreja, evitando vínculos cruzados acidentais.",
        "Imagens e PDFs não recebem URL pública fixa. Quando o armazenamento gerenciado é utilizado, o acesso é feito por URL assinada e temporária; o tempo padrão é de cinco minutos.",
        "A API valida tipos e limites de envio. Imagens de envelopes aceitam somente JPEG ou PNG e possuem limite de 4 MB antes de serem persistidas.",
        "As configurações de produção são validadas no início da aplicação: origem web, banco de dados, proxy, limite de corpo, credenciais e métricas não dependem de valores implícitos de desenvolvimento.",
    ])
    doc.add_heading("Controles de lançamento obrigatórios", level=2)
    add_body(doc, "Além das proteções já presentes na arquitetura, a liberação para produção deve ser condicionada aos controles operacionais abaixo. Eles são requisitos de implantação e não devem ser confundidos com funcionalidades de interface.")
    add_matrix(doc, ["Controle", "Critério de liberação"], [
        ("Proteção contra tentativa de login", "Limite por origem e por conta, resposta não enumerável e monitoramento de tentativas abusivas."),
        ("Logs e monitoramento", "Logs estruturados com mascaramento de senhas, tokens e dados pessoais; métricas e alertas sem exposição de PII."),
        ("Backup e restauração", "Rotina documentada e teste de restauração para banco de dados, imagens e PDFs privados."),
        ("Gestão de segredos", "Credenciais somente no ambiente de produção, com rotação, acesso mínimo e nenhuma chave exposta no código ou navegador."),
        ("Revisão de segurança", "Teste de fluxo com usuários reais, verificação de permissões, dependências e validação independente antes do rollout amplo."),
        ("Resposta a incidentes", "Responsáveis, canal de comunicação, prazo de contenção e procedimento de notificação definidos contratualmente."),
    ], [2700, 6660])
    add_callout(doc, "Transparência de segurança", "Funcionalidades como auditoria imutável de domínio, limitação completa de tentativas de login, logs com redação de dados sensíveis, métricas protegidas e restauração testada fazem parte do endurecimento de produção. Elas devem ser validadas no aceite de lançamento antes da expansão para toda a rede.", tone="amber")
    doc.add_heading("Arquitetura da solução", level=1)
    add_body(doc, "A solução utiliza uma arquitetura de monólito modular em monorepo. Essa escolha reduz a complexidade operacional inicial, mantendo fronteiras claras para evolução futura.")
    add_matrix(doc, ["Componente", "Responsabilidade"], [
        ("Aplicação web", "Interface administrativa responsiva, localizada e acessível pelo navegador."),
        ("API", "Regras de negócio, autenticação, autorização, validações, integração com banco e armazenamento privado."),
        ("Worker", "Base para processamento assíncrono e tarefas futuras de longa duração."),
        ("Banco PostgreSQL", "Dados administrativos, igrejas, membros, contribuições, sessões e metadados de arquivos."),
        ("Armazenamento de objetos", "Imagens de envelopes e PDFs, usando ambiente local no desenvolvimento e buckets privados na produção."),
    ], [2450, 6910])
    doc.add_heading("Tecnologias", level=2)
    add_bullets(doc, [
        "Next.js para a aplicação web; NestJS para a API; TypeScript de ponta a ponta.",
        "PostgreSQL e Drizzle ORM para modelagem e acesso aos dados.",
        "Supabase como opção de banco gerenciado e armazenamento privado em produção.",
        "Vercel como opção de hospedagem da aplicação web e API; arquitetura compatível com outros provedores gerenciados.",
        "Testes automatizados com Vitest e Playwright, incluindo validação visual em navegadores e tamanhos de tela suportados.",
    ])

    page_break(doc)
    add_kicker(doc, "07 | Experiência, qualidade e implantação")
    doc.add_heading("Experiência de uso", level=1)
    add_bullets(doc, [
        "Interface consistente em português do Brasil, inglês e espanhol, com seleção de idioma persistente no navegador.",
        "Navegação adaptada a desktop, largura estreita e dispositivos móveis, com controles adequados para toque.",
        "Mensagens de carregamento, sucesso, erro e estados vazios nas principais telas operacionais.",
        "Formatação de datas, números e valores monetários pela biblioteca internacional do navegador, de acordo com o idioma da interface.",
        "Fluxos principais disponíveis para login, seleção de igreja, membros, envelopes e relatórios.",
    ])
    doc.add_heading("Qualidade de engenharia", level=1)
    add_body(doc, "O projeto possui uma cadeia de qualidade que inclui formatação, lint, checagem de tipos, testes unitários, testes de integração, build de produção e testes de navegador. A validação visual complementa os testes funcionais para identificar textos cortados, controles menores que o esperado e rolagem horizontal indevida.")
    add_matrix(doc, ["Verificação", "Objetivo"], [
        ("Formatação e lint", "Padronizar o código e prevenir problemas comuns antes da publicação."),
        ("Tipos e contratos", "Garantir que web, API e banco conversem com formatos compatíveis e validados."),
        ("Testes automatizados", "Verificar autenticação, permissões, segregação entre igrejas, serviços e regras de domínio."),
        ("Testes de navegador", "Validar os fluxos principais da interface em ambiente próximo do uso real."),
        ("Regressão visual", "Manter acabamento profissional nos idiomas e larguras de tela suportados."),
    ], [2700, 6660])
    doc.add_heading("Implantação recomendada", level=1)
    add_numbered(doc, [
        "Criar o ambiente de produção nos Estados Unidos, com banco de dados e armazenamento privado na região definida pela organização.",
        "Configurar domínios, HTTPS, variáveis seguras, credenciais separadas de execução e de migração, e política de acesso à infraestrutura.",
        "Aplicar as migrações do banco e criar o primeiro administrador da plataforma e a primeira igreja piloto.",
        "Realizar importação inicial de dados, quando houver, e treinamento dos responsáveis da igreja piloto.",
        "Executar aceite funcional e visual do piloto antes do rollout progressivo para as demais unidades.",
        "Estabelecer monitoramento, rotina de backup, procedimento de restauração e suporte operacional antes da expansão completa.",
    ])
    add_callout(doc, "Infraestrutura", "O custo de hospedagem, banco, armazenamento, domínio, monitoramento e eventuais serviços de e-mail deve ser tratado separadamente da licença ou do desenvolvimento da plataforma. O ambiente deve ficar em conta corporativa do cliente ou em conta gerenciada conforme o contrato.", tone="amber")

    page_break(doc)
    add_kicker(doc, "08 | Escopo, evolução e premissas")
    doc.add_heading("Escopo principal da versão de lançamento", level=1)
    add_bullets(doc, [
        "Autenticação administrativa, seleção de igreja e permissões por função.",
        "Gestão segura de igrejas pela administração central.",
        "Cadastro e manutenção de membros por igreja.",
        "Registro manual de contribuições identificadas ou anônimas, com imagens privadas de envelopes.",
        "Histórico de contribuições por período e membro, com totais consolidados.",
        "Geração, armazenamento e download de relatórios PDF privados.",
        "Interface responsiva em português, inglês e espanhol.",
        "Banco PostgreSQL e armazenamento privado preparados para implantação gerenciada nos EUA.",
    ])
    doc.add_heading("Evoluções recomendadas", level=1)
    add_matrix(doc, ["Evolução", "Finalidade"], [
        ("Auditoria imutável", "Registrar alterações de membros, contribuições, relatórios e eventos de acesso com ator, data e alvo."),
        ("Fundos e formas configuráveis", "Permitir categorias como dízimo, ofertas, missões e outras destinações, com ciclo ativo/arquivado."),
        ("Lançamentos em lote", "Acelerar a contagem e a digitação de múltiplos envelopes com total de controle."),
        ("Correções por estorno", "Preservar o histórico financeiro ao corrigir um lançamento sem editar ou apagar o fato original."),
        ("CSV e exportações assíncronas", "Atender integrações contábeis e relatórios extensos sem bloquear a operação diária."),
        ("Doação digital", "Integrar pagamentos online, recorrência, conciliação e taxas de processamento após definição operacional e jurídica."),
        ("Operação avançada", "Backups testados, recuperação, observabilidade, alertas e metas formais de disponibilidade."),
    ], [2800, 6560])
    doc.add_heading("Fora do escopo inicial", level=2)
    add_bullets(doc, [
        "Aplicativos móveis nativos para iOS e Android; a prioridade é a aplicação web responsiva.",
        "Contabilidade geral, folha de pagamento, despesas, orçamento e demais funções de ERP.",
        "Processamento de pagamentos online sem definição prévia de parceiro, fluxo financeiro, exigências legais e taxas.",
        "Exclusão física rotineira de membros, igrejas ou contribuições; a diretriz é arquivar, inativar ou estornar, preservando histórico.",
    ])
    add_callout(doc, "Premissa contratual", "Este dossiê descreve a solução, seu escopo funcional e sua direção de implantação. Datas, integrações, regras específicas de operação, níveis de suporte, propriedade intelectual, custos de terceiros e critérios de aceite devem ser definidos em proposta comercial e contrato próprios.", tone="amber")

    page_break(doc)
    add_kicker(doc, "09 | Próximos passos")
    doc.add_heading("Roteiro de contratação e implantação", level=1)
    add_numbered(doc, [
        "Validação do escopo funcional, perfis de usuário, dados iniciais e indicadores que a administração central precisa acompanhar.",
        "Definição do modelo comercial: licença corporativa, desenvolvimento sob encomenda ou compra de código-fonte.",
        "Definição da infraestrutura, responsáveis técnicos, domínio e política de acesso ao ambiente de produção.",
        "Configuração de ambiente, carga piloto e treinamento da primeira igreja.",
        "Aceite do piloto com usuários reais e ajustes priorizados pela operação.",
        "Expansão gradativa para as demais igrejas, acompanhada por suporte e monitoramento.",
    ])
    doc.add_heading("Informações a confirmar com o cliente", level=2)
    add_matrix(doc, ["Tema", "Decisão necessária"], [
        ("Rollout", "Quantas igrejas entram no piloto e qual é a meta de expansão por mês?"),
        ("Dados existentes", "Há planilhas, sistemas anteriores ou imagens que precisam ser migrados?"),
        ("Governança", "Quem aprova usuários, acessos, arquivamento de igrejas e relatórios financeiros?"),
        ("Relatórios", "Quais formatos, periodicidades, moedas, fusos e campos precisam ser obrigatórios?"),
        ("Segurança", "Quais políticas internas, retenções, auditorias e exigências de compliance são aplicáveis?"),
        ("Suporte", "Quais horários, canais, tempos de resposta e responsáveis devem constar no acordo operacional?"),
    ], [2500, 6860])
    doc.add_heading("Conclusão", level=1)
    add_body(doc, "O UCKG Donations foi concebido como uma base segura e profissional para padronizar a gestão de contribuições em uma rede de igrejas. A solução combina operação local por igreja, governança central, controle de acesso, documentação de envelopes e relatórios privados, com uma arquitetura preparada para evolução progressiva.")
    add_body(doc, "A fase de implantação deve confirmar os processos internos, validar a operação piloto e consolidar os requisitos de governança antes do rollout para toda a rede.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UCKG Donations | Clareza para cuidar. Segurança para servir.")
    set_run_font(r, 10.5, NAVY, bold=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = "UCKG Donations - Dossiê funcional e técnico"
    doc.core_properties.subject = "Apresentação de produto e especificações"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
